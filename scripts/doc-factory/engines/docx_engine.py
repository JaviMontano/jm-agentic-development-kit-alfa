"""DOCX Engine — Word document generation via docxtpl (Jinja2 + python-docx)."""

from pathlib import Path

from generate import register_engine


@register_engine('docx')
class DocxEngine:
    def generate(self, data, template_path, output_path, schema=None):
        """
        Generate DOCX from template + data.

        If template_path is a real .docx with {placeholders}, uses docxtpl.
        If no template, falls back to python-docx programmatic generation.
        """
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        if template_path and Path(template_path).exists() and str(template_path).endswith('.docx'):
            self._from_template(data, template_path, output_path)
        else:
            self._programmatic(data, output_path)

    def _from_template(self, data, template_path, output_path):
        """Fill a .docx template using docxtpl (Jinja2 syntax)."""
        try:
            from docxtpl import DocxTemplate
        except ImportError:
            raise RuntimeError("docxtpl required: pip install docxtpl")

        doc = DocxTemplate(str(template_path))

        # Compute totals if items exist
        if 'items' in data and 'total' not in data:
            data['total'] = sum(
                i.get('quantity', 0) * i.get('unit_price', 0)
                for i in data['items']
            )
        if 'items' in data and 'subtotal' not in data:
            data['subtotal'] = data.get('total', 0)
            tax_rate = data.get('tax_rate', 0)
            data['tax_amount'] = data['subtotal'] * tax_rate
            data['grand_total'] = data['subtotal'] + data['tax_amount']

        from datetime import datetime
        data.setdefault('_generated_at', datetime.now().strftime('%Y-%m-%d %H:%M'))

        doc.render(data)
        doc.save(str(output_path))

    def _programmatic(self, data, output_path):
        """Generate DOCX programmatically when no template is available."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError("python-docx required: pip install python-docx")

        doc = Document()

        # Style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Header with folio
        if data.get('folio'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(data['folio'])
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)

        # Title
        title = data.get('title', data.get('name', 'Document'))
        heading = doc.add_heading(title, level=1)

        # Metadata
        meta_parts = []
        if data.get('date'):
            meta_parts.append(f"Fecha: {data['date']}")
        if data.get('client_name'):
            meta_parts.append(f"Para: {data['client_name']}")
        if meta_parts:
            doc.add_paragraph(' | '.join(meta_parts))

        # Sections
        for section in data.get('sections', []):
            doc.add_heading(section.get('title', ''), level=2)
            doc.add_paragraph(section.get('content', ''))

        # Items table
        if data.get('items'):
            doc.add_heading('Detalle', level=2)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            headers = ['#', 'Descripcion', 'Cantidad', 'Precio', 'Total']
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for paragraph in cell.paragraphs:
                    paragraph.runs[0].font.bold = True

            for idx, item in enumerate(data['items'], 1):
                row = table.add_row()
                qty = item.get('quantity', 0)
                price = item.get('unit_price', 0)
                row.cells[0].text = str(idx)
                row.cells[1].text = item.get('description', '')
                row.cells[2].text = str(qty)
                row.cells[3].text = f"{price:.2f}"
                row.cells[4].text = f"{qty * price:.2f}"

        # Total
        if 'total' in data or data.get('items'):
            total = data.get('total', sum(i.get('quantity', 0) * i.get('unit_price', 0) for i in data.get('items', [])))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"Total: {data.get('currency', '$')} {total:.2f}")
            run.font.size = Pt(14)
            run.font.bold = True

        # Content block
        if data.get('content'):
            doc.add_paragraph(data['content'])

        # Footer
        from datetime import datetime
        doc.add_paragraph('')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{'Documento ' + data['folio'] + ' | ' if data.get('folio') else ''}Generado {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.save(str(output_path))
