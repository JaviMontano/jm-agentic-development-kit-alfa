#!/usr/bin/env python3
"""
Pre-Generation Validator (SDD Gate G0 + G1)
============================================
Validates inputs BEFORE document generation.

G0: Schema validation — data matches JSON Schema contract
G1: Consistency — template placeholders match schema fields

Usage:
    python validate.py --schema schema.json --data data.json
    python validate.py --schema schema.json --data data.json --template template.docx
"""

import argparse
import json
import re
import sys
from pathlib import Path


def load_json(path):
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def gate_g0_schema(schema, data):
    """Gate G0: Validate data against JSON Schema (required fields, types)."""
    errors = []

    required = schema.get('required', [])
    properties = schema.get('properties', {})

    # Check required fields
    for field in required:
        if field not in data:
            errors.append(f"G0: Missing required field: '{field}'")

    # Check types for present fields
    for field, value in data.items():
        if field not in properties:
            continue  # Extra fields are OK (additionalProperties)

        expected_type = properties[field].get('type')
        if expected_type:
            type_map = {
                'string': str, 'number': (int, float), 'integer': int,
                'boolean': bool, 'array': list, 'object': dict
            }
            expected_python = type_map.get(expected_type)
            if expected_python and not isinstance(value, expected_python):
                errors.append(f"G0: Field '{field}' expected {expected_type}, got {type(value).__name__}")

        # Check pattern (regex)
        pattern = properties[field].get('pattern')
        if pattern and isinstance(value, str):
            if not re.match(pattern, value):
                errors.append(f"G0: Field '{field}' does not match pattern: {pattern}")

        # Check enum
        enum_values = properties[field].get('enum')
        if enum_values and value not in enum_values:
            errors.append(f"G0: Field '{field}' must be one of {enum_values}, got '{value}'")

        # Check minimum
        minimum = properties[field].get('minimum')
        if minimum is not None and isinstance(value, (int, float)) and value < minimum:
            errors.append(f"G0: Field '{field}' must be >= {minimum}, got {value}")

        # Check minLength
        min_length = properties[field].get('minLength')
        if min_length and isinstance(value, str) and len(value) < min_length:
            errors.append(f"G0: Field '{field}' must be at least {min_length} chars")

        # Check minItems for arrays
        min_items = properties[field].get('minItems')
        if min_items and isinstance(value, list) and len(value) < min_items:
            errors.append(f"G0: Field '{field}' must have at least {min_items} items")

    return errors


def gate_g1_consistency(schema, data, template_path):
    """Gate G1: Validate template-schema-data consistency."""
    errors = []

    if not template_path or not Path(template_path).exists():
        return errors  # No template to validate against

    template_ext = Path(template_path).suffix.lower()

    # Extract placeholders from template
    placeholders = set()

    if template_ext in ('.html', '.md', '.txt'):
        content = Path(template_path).read_text(encoding='utf-8')
        # Jinja2 / Handlebars: {{ var }}, {% for %}, {# comment #}
        placeholders = set(re.findall(r'\{\{[\s]*(\w+)', content))
        # Also find loop variables: {% for x in items %}
        placeholders |= set(re.findall(r'\{%\s*for\s+\w+\s+in\s+(\w+)', content))

    elif template_ext == '.docx':
        try:
            # docxtpl uses {var} syntax
            from docx import Document
            doc = Document(str(template_path))
            text = '\n'.join(p.text for p in doc.paragraphs)
            text += '\n'.join(cell.text for t in doc.tables for row in t.rows for cell in row.cells)
            placeholders = set(re.findall(r'\{[\s]*(\w+)', text))
            # Remove common false positives
            placeholders -= {'d', 'endif', 'endfor', 'else', 'if', 'for'}
        except Exception:
            pass  # Can't read template — skip G1

    if placeholders:
        schema_fields = set(schema.get('properties', {}).keys())
        data_fields = set(data.keys())

        # Placeholders in template but not in data
        missing_in_data = placeholders - data_fields - {'loop', '_generated_at'}
        if missing_in_data:
            errors.append(f"G1: Template has placeholders not in data: {missing_in_data}")

    return errors


def validate_inputs(schema_path, data_path, template_path=None, formats=None):
    """Run all validation gates. Returns True if all pass."""
    errors = []

    # Load data
    try:
        data = load_json(data_path)
    except Exception as e:
        print(f"FAIL: Cannot load data file: {e}")
        return False

    # G0: Schema validation
    if schema_path:
        try:
            schema = load_json(schema_path)
            g0_errors = gate_g0_schema(schema, data)
            errors.extend(g0_errors)
        except Exception as e:
            errors.append(f"G0: Cannot load schema: {e}")
    else:
        schema = {}

    # G1: Consistency
    if template_path:
        g1_errors = gate_g1_consistency(schema, data, template_path)
        errors.extend(g1_errors)

    # Report
    if errors:
        print(f"VALIDATION FAILED: {len(errors)} error(s)")
        for err in errors:
            print(f"  ✗ {err}")
        return False

    print("VALIDATION PASSED: All gates clear")
    return True


def main():
    parser = argparse.ArgumentParser(description='Pre-Generation Validator (Gates G0 + G1)')
    parser.add_argument('--schema', help='Path to JSON Schema')
    parser.add_argument('--data', required=True, help='Path to data JSON')
    parser.add_argument('--template', help='Path to template file')

    args = parser.parse_args()
    ok = validate_inputs(args.schema, args.data, args.template)
    sys.exit(0 if ok else 1)


if __name__ == '__main__':
    main()
